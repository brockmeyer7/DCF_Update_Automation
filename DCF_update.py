import os
from docx import Document
from docx.shared import RGBColor, Pt, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from datetime import datetime
from dateutil import parser


def get_docx(path: str) -> list:
    results = []
    try:
        for fname in os.listdir(path):
            if fname.endswith('.docx'):
                results.append(path + '/' + fname)
    except:
        pass
    return results


def find_DCFs(comparison_date, files: list) -> list:
    results = []
    for f in files:
        try:
            document = Document(f)
            section = document.sections[0]
            header = section.header
            table = document.tables[0]
            contents = table.cell(2, 1).text
            d = parser.parse(contents)
            if d >= comparison_date and header.paragraphs[1].text.endswith('Rev. D'):
                results.append(f)
        except:
            pass
    return results


def update_DCF(file_path):

    doc = Document(file_path)

    # Update header
    h = doc.sections[0].header
    p = h.paragraphs
    p[0].text = 'Design Control Form'
    p[1].text = 'QAF 080303  Rev. E'
    p[1].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p = p[0].insert_paragraph_before('')
    r = p.add_run()
    r.add_picture('ESP_LOGO.png')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_p = h.add_paragraph()
    new_p.paragraph_format.line_spacing = Pt(6)

    # Add section 9
    run = doc.add_paragraph().add_run('SECTION 9: Revision History\n')
    run.font.color.rgb = RGBColor(0, 0, 128)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True

    # Add revision table
    table = doc.add_table(rows=3, cols=4, style="Table Grid")

    # Set table row heights
    table.rows[0].height = 502900
    table.rows[1].height = 274320
    table.rows[2].height = 274320

    # Set table column widths
    for c in table.columns[0].cells:
        c.width = Emu(411480)
    for c in table.columns[1].cells:
        c.width = Emu(731520)
    for c in table.columns[2].cells:
        c.width = Emu(4206240)
    for c in table.columns[3].cells:
        c.width = Emu(932688)

    # Add coumn header text
    table.cell(0, 0).paragraphs[0].add_run('Rev')
    table.cell(0, 1).paragraphs[0].add_run('Release Date')
    table.cell(0, 2).paragraphs[0].add_run('Revision Description')
    table.cell(0, 3).paragraphs[0].add_run('Author')

    # Add revision text
    table.cell(1, 0).paragraphs[0].add_run('E')
    table.cell(1, 1).paragraphs[0].add_run('1-03-22')
    table.cell(1, 2).paragraphs[0].add_run('Format and minor wording updates.')
    table.cell(1, 3).paragraphs[0].add_run('D. Crow')

    # Update text formatting
    for c in table.rows[0].cells:
        p = c.paragraphs[0]
        f = c.paragraphs[0].runs[0].font
        f.name = 'Arial'
        f.size = Pt(11)
        f.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for c in table.rows[1].cells:
        p = c.paragraphs[0]
        f = c.paragraphs[0].runs[0].font
        f.name = 'Arial'
        f.size = Pt(10)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Save edits
    doc.save(file_path)


if __name__ == '__main__':
    comparison_date = datetime.strptime('01/01/2022', '%m/%d/%Y')
    rootdir = 'U:/Product/PRODUCT SPECIFICATIONS'
    dirs = [rootdir + '/' + i for i in next(os.walk(rootdir))[1]]
    files_to_update = []
    dirs_to_update = []

    for path in dirs:
        docx_files = get_docx(path)
        if len(docx_files) > 0:
            files_to_update.extend(find_DCFs(comparison_date, docx_files))

    for file in files_to_update:
        file_path = file
        directory = file_path.rsplit('/', 1)[0]
        update_DCF(file_path)

    with open('C:/Users/bmeyer/Documents/DCF REV FIX/Updated_Files.txt', 'w') as f:
        for i, filename in enumerate(files_to_update):
            f.write(str(i) + ': ' + filename + '\n')
